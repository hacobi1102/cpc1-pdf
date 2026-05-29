# -*- coding: utf-8 -*-
"""
app.py – Trạm Xử Lý PDF Chuyên Nghiệp v3.0 (Web App)
Backend Flask – xử lý PDF: gộp, tách, xuất ảnh, chuyển Word (OCR)
"""

import os
import sys
import io
import uuid
import zipfile
import atexit
import shutil

from flask import Flask, request, jsonify, send_file, render_template

from pypdf import PdfReader, PdfWriter

# ─── Thư viện tuỳ chọn ───
HAS_FITZ = True
HAS_OCR  = True

try:
    import fitz
    from PIL import Image
except ImportError:
    HAS_FITZ = False

try:
    import pytesseract
except ImportError:
    HAS_OCR = False

# ─── Auto-config Tesseract ───
def _auto_tesseract():
    if not HAS_OCR:
        return
    from shutil import which
    for t in [r"C:\Program Files\Tesseract-OCR\tesseract.exe",
              r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
              r"C:\Tesseract-OCR\tesseract.exe"]:
        if os.path.isfile(t):
            pytesseract.pytesseract.tesseract_cmd = t
            return
    found = which("tesseract")
    if found:
        pytesseract.pytesseract.tesseract_cmd = found

_auto_tesseract()

# ─── Flask App ───
app = Flask(__name__)
app.config['TEMPLATES_AUTO_RELOAD'] = True

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "temp_uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Dọn dẹp thư mục tạm khi tắt server
atexit.register(lambda: shutil.rmtree(UPLOAD_DIR, ignore_errors=True))

# Registry lưu thông tin file đã upload  { file_id -> {path, name, total_pages} }
files_db: dict[str, dict] = {}


# ─── Routes ───

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/upload", methods=["POST"])
def upload():
    """Upload một hoặc nhiều file PDF."""
    results = []
    for f in request.files.getlist("files"):
        if not f.filename or not f.filename.lower().endswith(".pdf"):
            continue
        file_id = uuid.uuid4().hex[:10]
        path = os.path.join(UPLOAD_DIR, f"{file_id}.pdf")
        f.save(path)
        try:
            reader = PdfReader(path)
            total  = len(reader.pages)
        except Exception as e:
            os.remove(path)
            return jsonify({"error": f"Lỗi đọc {f.filename}: {e}"}), 400
        files_db[file_id] = {
            "path":        path,
            "name":        f.filename,
            "total_pages": total,
        }
        results.append({
            "file_id":     file_id,
            "name":        f.filename,
            "total_pages": total,
        })
    return jsonify(results)


@app.route("/api/files/<file_id>", methods=["DELETE"])
def delete_file(file_id):
    """Xóa một file khỏi danh sách."""
    info = files_db.pop(file_id, None)
    if info:
        try:
            os.remove(info["path"])
        except OSError:
            pass
    return jsonify({"ok": True})


@app.route("/api/files/clear", methods=["DELETE"])
def clear_files():
    """Xóa tất cả file."""
    for fid, info in list(files_db.items()):
        try:
            os.remove(info["path"])
        except OSError:
            pass
    files_db.clear()
    return jsonify({"ok": True})


@app.route("/api/thumb/<file_id>/<int:page>")
def get_thumb(file_id, page):
    """Trả thumbnail PNG cho một trang PDF."""
    if not HAS_FITZ:
        return "PyMuPDF not installed", 500
    info = files_db.get(file_id)
    if not info:
        return "File not found", 404
    doc = fitz.open(info["path"])
    if page < 0 or page >= len(doc):
        doc.close()
        return "Page not found", 404
    p = doc[page]
    scale = min(280 / p.rect.width, 360 / p.rect.height)
    mat   = fitz.Matrix(scale, scale)
    pix   = p.get_pixmap(matrix=mat, alpha=False)
    data  = pix.tobytes("png")
    doc.close()
    return send_file(io.BytesIO(data), mimetype="image/png",
                     download_name=f"thumb_{file_id}_{page}.png")


def compress_pdf_buffer(buf):
    if not HAS_FITZ:
        return buf
    try:
        doc = fitz.open(stream=buf.getvalue(), filetype="pdf")
        out_buf = io.BytesIO()
        doc.save(out_buf, garbage=4, deflate=True, clean=True)
        doc.close()
        out_buf.seek(0)
        return out_buf
    except Exception:
        buf.seek(0)
        return buf


@app.route("/api/merge", methods=["POST"])
def merge():
    """Gộp các trang đã chọn từ nhiều file thành 1 PDF."""
    data = request.json  # {files: [{file_id, pages: [0,1,...]}], compress: bool}
    compress = data.get("compress", False)
    writer = PdfWriter()
    for item in data["files"]:
        info = files_db.get(item["file_id"])
        if not info:
            return jsonify({"error": f"File {item['file_id']} not found"}), 404
        reader = PdfReader(info["path"])
        for i in item["pages"]:
            if 0 <= i < len(reader.pages):
                writer.add_page(reader.pages[i])
    buf = io.BytesIO()
    writer.write(buf)
    if compress and HAS_FITZ:
        buf = compress_pdf_buffer(buf)
    else:
        buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="merged.pdf",
                     mimetype="application/pdf")


@app.route("/api/split", methods=["POST"])
def split():
    """Tách các trang đã chọn thành file riêng, trả ZIP."""
    data = request.json  # {file_id, pages: [0,1,...], compress: bool}
    compress = data.get("compress", False)
    info = files_db.get(data["file_id"])
    if not info:
        return jsonify({"error": "File not found"}), 404
    reader = PdfReader(info["path"])
    base   = os.path.splitext(info["name"])[0]
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i in data["pages"]:
            if 0 <= i < len(reader.pages):
                w = PdfWriter()
                w.add_page(reader.pages[i])
                pbuf = io.BytesIO()
                w.write(pbuf)
                if compress and HAS_FITZ:
                    pbuf = compress_pdf_buffer(pbuf)
                else:
                    pbuf.seek(0)
                zf.writestr(f"{base}_trang_{i+1}.pdf", pbuf.getvalue())
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name=f"{base}_tach.zip",
                     mimetype="application/zip")


@app.route("/api/compress", methods=["POST"])
def do_compress_endpoint():
    """Nén các trang đã chọn thành file mới."""
    data = request.json
    info = files_db.get(data["file_id"])
    if not info:
        return jsonify({"error": "File not found"}), 404
    reader = PdfReader(info["path"])
    base = os.path.splitext(info["name"])[0]
    
    writer = PdfWriter()
    for i in data["pages"]:
        if 0 <= i < len(reader.pages):
            writer.add_page(reader.pages[i])
            
    buf = io.BytesIO()
    writer.write(buf)
    
    if HAS_FITZ:
        buf = compress_pdf_buffer(buf)
    else:
        buf.seek(0)
        
    return send_file(buf, as_attachment=True,
                     download_name=f"{base}_nen.pdf",
                     mimetype="application/pdf")


@app.route("/api/images", methods=["POST"])
def to_images():
    """Xuất các trang đã chọn thành ảnh PNG, trả ZIP."""
    if not HAS_FITZ:
        return jsonify({"error": "Cần cài PyMuPDF: pip install PyMuPDF"}), 500
    data = request.json  # {file_id, pages: [0,1,...]}
    info = files_db.get(data["file_id"])
    if not info:
        return jsonify({"error": "File not found"}), 404
    doc  = fitz.open(info["path"])
    base = os.path.splitext(info["name"])[0]
    buf  = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i in data["pages"]:
            if 0 <= i < len(doc):
                pix = doc[i].get_pixmap(dpi=200, alpha=False)
                zf.writestr(f"{base}_trang_{i+1}.png", pix.tobytes("png"))
    doc.close()
    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name=f"{base}_anh.zip",
                     mimetype="application/zip")


@app.route("/api/word", methods=["POST"])
def to_word():
    """Chuyển các trang thành Word (.docx), hỗ trợ OCR."""
    if not HAS_FITZ:
        return jsonify({"error": "Cần cài PyMuPDF: pip install PyMuPDF"}), 500
    data = request.json
    info = files_db.get(data["file_id"])
    if not info:
        return jsonify({"error": "File not found"}), 404

    ocr_mode = data.get("ocr_mode")
    if ocr_mode is None:
        use_ocr = data.get("ocr", False)
        ocr_mode = "advanced" if use_ocr else "none"

    if ocr_mode in ["basic", "advanced"] and not HAS_OCR:
        return jsonify({"error": "Cần cài pytesseract: pip install pytesseract"}), 500

    base = os.path.splitext(info["name"])[0]
    buf = io.BytesIO()
    import tempfile

    if ocr_mode == "none":
        try:
            from pdf2docx import Converter
        except ImportError:
            return jsonify({"error": "Thiếu thư viện pdf2docx"}), 500
            
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_docx = tmp.name
        try:
            cv = Converter(info["path"])
            cv.convert(tmp_docx, pages=data["pages"])
            cv.close()
            with open(tmp_docx, "rb") as f:
                buf.write(f.read())
        except Exception as e:
            return jsonify({"error": f"Lỗi chuyển Word: {e}"}), 500
        finally:
            if os.path.exists(tmp_docx):
                try: os.remove(tmp_docx)
                except: pass

    elif ocr_mode == "basic":
        from docx import Document
        doc = fitz.open(info["path"])
        word_doc = Document()
        pages = data["pages"]
        for idx, i in enumerate(pages):
            if 0 <= i < len(doc):
                page = doc[i]
                pix = page.get_pixmap(dpi=200, alpha=False)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                text = pytesseract.image_to_string(img, lang="vie+eng")
                word_doc.add_paragraph(text)
                if idx < len(pages) - 1:
                    word_doc.add_page_break()
        doc.close()
        word_doc.save(buf)

    elif ocr_mode == "advanced":
        try:
            import ocrmypdf
            from pdf2docx import Converter
        except ImportError:
            return jsonify({"error": "Thiếu thư viện ocrmypdf hoặc pdf2docx"}), 500
            
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_in, \
             tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_out, \
             tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
            tmp_in_path = tmp_in.name
            tmp_out_path = tmp_out.name
            tmp_docx_path = tmp_docx.name

        try:
            doc = fitz.open(info["path"])
            doc.select(data["pages"])
            doc.save(tmp_in_path)
            doc.close()

            ocrmypdf.ocr(
                tmp_in_path, 
                tmp_out_path, 
                language="vie+eng", 
                deskew=True, 
                force_ocr=True, 
                optimize=1
            )

            # --- SỬA LỖI pdf2docx RA ẢNH ---
            # ocrmypdf tạo ra text ẩn (Render mode 3) và giữ nguyên ảnh nền.
            # pdf2docx mặc định bỏ qua text ẩn nên chỉ trích xuất được ảnh nền.
            # Ta cần xóa ảnh nền và chuyển text ẩn thành text hiện (0 Tr).
            doc_fix = fitz.open(tmp_out_path)
            for page in doc_fix:
                for item in page.get_images():
                    page.delete_image(item[0])
            
            for xref in range(1, doc_fix.xref_length()):
                stream = doc_fix.xref_stream(xref)
                if stream and b"3 Tr" in stream:
                    stream = stream.replace(b"3 Tr", b"0 Tr")
                    doc_fix.update_stream(xref, stream)
            
            fixed_pdf_path = tmp_out_path + "_fixed.pdf"
            doc_fix.save(fixed_pdf_path)
            doc_fix.close()
            # ---------------------------------

            cv = Converter(fixed_pdf_path)
            cv.convert(tmp_docx_path)
            cv.close()

            with open(tmp_docx_path, "rb") as f:
                buf.write(f.read())
        except Exception as e:
            print(f"OCRmyPDF failed: {e}. Fallback to basic OCR.")
            buf = io.BytesIO()
            from docx import Document
            doc = fitz.open(info["path"])
            word_doc = Document()
            for idx, i in enumerate(data["pages"]):
                if 0 <= i < len(doc):
                    page = doc[i]
                    pix = page.get_pixmap(dpi=200, alpha=False)
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    text = pytesseract.image_to_string(img, lang="vie+eng")
                    word_doc.add_paragraph(text)
                    if idx < len(data["pages"]) - 1:
                        word_doc.add_page_break()
            doc.close()
            word_doc.save(buf)
        finally:
            for path in [tmp_in_path, tmp_out_path, tmp_docx_path, tmp_out_path + "_fixed.pdf"]:
                if os.path.exists(path):
                    try: os.remove(path)
                    except: pass

    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name=f"{base}.docx",
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/api/excel", methods=["POST"])
def to_excel():
    """Trích xuất bảng từ PDF và chuyển sang Excel (.xlsx)."""
    try:
        import pdfplumber
        import pandas as pd
    except ImportError:
        return jsonify({"error": "Đang cài đặt thư viện hỗ trợ (pdfplumber, pandas). Vui lòng chờ..."}), 500

    data = request.json
    info = files_db.get(data["file_id"])
    if not info:
        return jsonify({"error": "File not found"}), 404

    pages = data["pages"]
    base = os.path.splitext(info["name"])[0]
    buf = io.BytesIO()
    
    try:
        with pdfplumber.open(info["path"]) as pdf:
            with pd.ExcelWriter(buf, engine='openpyxl') as writer:
                table_found = False
                for p_idx in pages:
                    if 0 <= p_idx < len(pdf.pages):
                        page = pdf.pages[p_idx]
                        tables = page.extract_tables()
                        for t_idx, table in enumerate(tables):
                            if not table:
                                continue
                            table_found = True
                            df = pd.DataFrame(table)
                            sheet_name = f"Trang {p_idx+1}"
                            if len(tables) > 1:
                                sheet_name += f" - Bảng {t_idx+1}"
                            sheet_name = sheet_name[:31]
                            df.to_excel(writer, sheet_name=sheet_name, index=False, header=False)
                
                if not table_found:
                    df = pd.DataFrame([["Không tìm thấy bảng biểu nào trên các trang đã chọn."]])
                    df.to_excel(writer, sheet_name="Thông báo", index=False, header=False)
    except Exception as e:
        return jsonify({"error": f"Lỗi xử lý Excel: {str(e)}"}), 500

    buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name=f"{base}.xlsx",
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@app.route("/api/status")
def status():
    """Trả thông tin trạng thái server."""
    return jsonify({
        "fitz":      HAS_FITZ,
        "ocr":       HAS_OCR,
        "files":     len(files_db),
    })


# ─── Khởi chạy ───
if __name__ == "__main__":
    import threading
    import webbrowser
    
    port_env = os.environ.get("PORT")
    port = int(port_env) if port_env else 5000
    
    if port_env is None:
        print(f"\n  PDF Toolkit Web đang chạy tại: http://localhost:{port}\n")
        threading.Timer(1.0, lambda: webbrowser.open(f"http://localhost:{port}")).start()
    else:
        print(f"\n  Production server starting on port {port}\n")
        
    app.run(host="0.0.0.0", port=port, debug=False)
